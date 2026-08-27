# 注意：本模块当前未接入主链路（独立文档解析适配器），修复保留待未来接线
"""
龙虾文档适配模块 (Lobster Document Adapter)
支持 PDF、DOCX、HTML 等格式文档的解析与上下文注入
用于将文档内容转换为 LLM 可理解的上下文格式
"""

import io
import os
import re
import zipfile
import logging
from typing import Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# ========== 解析防护上限（修复：原先无任何大小限制，可被超大文件/zip炸弹打爆内存） ==========

# 输入文档大小上限，默认 20MB，可用环境变量 LOBSTER_MAX_BYTES 覆盖
LOBSTER_MAX_BYTES = int(os.environ.get("LOBSTER_MAX_BYTES", 20 * 1024 * 1024))
# PDF 页数上限
LOBSTER_MAX_PDF_PAGES = 500
# docx(zip) 总未压缩体积 / 压缩体积 的膨胀比上限，超过视为 zip 炸弹
LOBSTER_ZIP_RATIO_LIMIT = 100
# docx(zip) 内部条目数上限
LOBSTER_ZIP_MAX_ENTRIES = 200


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    title: str = ""
    content: str = ""
    sections: list = field(default_factory=list)  # [{heading, content, level}]
    tables: list = field(default_factory=list)  # [markdown_table_str]
    metadata: dict = field(default_factory=dict)  # {author, created, pages, etc.}
    source_type: str = ""  # pdf, docx, html
    token_estimate: int = 0

    def to_context_string(self, max_tokens: int = 8000) -> str:
        """将文档转换为适合注入LLM上下文的字符串"""
        # Token估算: 约4字符=1token
        # 如果总内容超过max_tokens,进行智能截断
        max_chars = max_tokens * 3  # conservative: 3 chars per token

        if not self.content:
            return ""

        # 构建上下文字符串
        parts = []

        # 文档元信息
        header = f"[Document: {self.title}]"
        if self.source_type:
            header += f" (type: {self.source_type})"
        if self.metadata:
            meta_items = [f"{k}: {v}" for k, v in self.metadata.items() if v]
            if meta_items:
                header += "\n" + " | ".join(meta_items)
        parts.append(header)

        # 主体内容
        content = self.content

        # 如果有表格,附加到内容后
        if self.tables:
            content += "\n\n--- Tables ---\n\n" + "\n\n".join(self.tables)

        parts.append(content)

        result = "\n\n".join(parts)

        # 智能截断
        if len(result) > max_chars:
            result = result[:max_chars]
            # 尝试在最后一个完整句子处截断
            last_period = result.rfind("。")
            last_period_en = result.rfind(". ")
            cut_pos = max(last_period, last_period_en)
            if cut_pos > max_chars * 0.8:
                result = result[:cut_pos + 1]
            result += "\n\n[... 文档已截断 ...]"

        return result

    def to_openai_messages(self, system_prompt: str = None) -> list:
        """将文档转换为OpenAI消息格式,适合注入对话上下文"""
        messages = []

        if system_prompt:
            messages.append({
                "role": "system",
                "content": system_prompt
            })

        context = self.to_context_string()
        if context:
            messages.append({
                "role": "system",
                "content": f"以下是供你参考的文档内容:\n\n{context}"
            })

        return messages


class LobsterDocAdapter:
    """龙虾文档适配器 - 多格式文档解析与上下文注入"""

    def __init__(self):
        self._pdf_available = False
        self._docx_available = False
        self._bs4_available = False
        self._check_dependencies()

    def _check_dependencies(self):
        """检查可选依赖是否可用"""
        try:
            import pypdf
            self._pdf_available = True
        except ImportError:
            logger.warning("pypdf未安装,PDF解析不可用. pip install pypdf")
        try:
            from docx import Document
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx未安装,DOCX解析不可用. pip install python-docx")
        try:
            from bs4 import BeautifulSoup
            self._bs4_available = True
        except ImportError:
            logger.warning("beautifulsoup4未安装,HTML解析不可用. pip install beautifulsoup4")

    @staticmethod
    def _check_size(content: bytes):
        """读文件前的统一大小检查（修复：原先无限制）"""
        if content is not None and len(content) > LOBSTER_MAX_BYTES:
            raise ValueError(
                f"文档过大: {len(content)} 字节，超过上限 "
                f"{LOBSTER_MAX_BYTES} 字节（约{LOBSTER_MAX_BYTES // (1024 * 1024)}MB），拒绝解析"
            )

    def parse(self, content: bytes, filename: str = "") -> ParsedDocument:
        """自动检测文件类型并解析"""
        ext = self._detect_extension(filename, content)
        if ext == "pdf":
            return self.parse_pdf(content)
        elif ext in ("docx", "doc"):
            return self.parse_docx(content)
        elif ext in ("html", "htm"):
            return self.parse_html(content)
        elif ext == "txt":
            return self.parse_text(content)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def parse_pdf(self, content: bytes) -> ParsedDocument:
        """解析PDF文档"""
        self._check_size(content)  # 修复：读文件前检查大小上限
        if not self._pdf_available:
            return ParsedDocument(source_type="pdf", content="[PDF解析不可用: pypdf未安装]")

        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        # 修复：PDF 页数上限，防止全页解压把内存/CPU 打爆
        if len(reader.pages) > LOBSTER_MAX_PDF_PAGES:
            raise ValueError(
                f"PDF页数过多: {len(reader.pages)} 页，超过上限 {LOBSTER_MAX_PDF_PAGES} 页，拒绝解析"
            )
        doc = ParsedDocument(source_type="pdf")
        doc.metadata = {
            "pages": len(reader.pages),
            "author": reader.metadata.author if reader.metadata else "",
            "title": reader.metadata.title if reader.metadata else "",
            "created": str(reader.metadata.creation_date) if reader.metadata and reader.metadata.creation_date else "",
        }

        full_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                full_text.append(text)

        doc.content = "\n\n".join(full_text)
        doc.title = doc.metadata.get("title") or "Untitled PDF"
        doc._estimate_tokens()
        return doc

    def parse_docx(self, content: bytes) -> ParsedDocument:
        """解析DOCX文档"""
        self._check_size(content)  # 修复：读文件前检查大小上限
        if not self._docx_available:
            return ParsedDocument(source_type="docx", content="[DOCX解析不可用: python-docx未安装]")

        # 修复：docx 本质是 zip，解压前先做 zip 炸弹检查（膨胀比/条目数）
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                infos = zf.infolist()
                if len(infos) > LOBSTER_ZIP_MAX_ENTRIES:
                    raise ValueError(
                        f"DOCX内部条目数过多: {len(infos)} 个，超过上限 "
                        f"{LOBSTER_ZIP_MAX_ENTRIES} 个，疑似zip炸弹，拒绝解析"
                    )
                total_uncompressed = sum(i.file_size for i in infos)
                total_compressed = sum(i.compress_size for i in infos)
                if total_compressed > 0 and \
                        total_uncompressed / total_compressed > LOBSTER_ZIP_RATIO_LIMIT:
                    raise ValueError(
                        f"DOCX解压膨胀比过高: 未压缩{total_uncompressed}/压缩{total_compressed}"
                        f"（比值>{LOBSTER_ZIP_RATIO_LIMIT}），疑似zip炸弹，拒绝解析"
                    )
        except zipfile.BadZipFile:
            # 非 zip 结构：交给 python-docx 自行报错
            pass

        from docx import Document

        doc_file = Document(io.BytesIO(content))
        doc = ParsedDocument(source_type="docx")

        # Extract sections by heading level
        current_section = {"heading": "Document Start", "content": [], "level": 0}

        for para in doc_file.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                # Save previous section
                if current_section["content"]:
                    current_section["content"] = "\n".join(current_section["content"])
                    doc.sections.append(current_section)

                level = int(para.style.name.replace("Heading ", "").replace("Heading", "1") or "1")
                current_section = {
                    "heading": para.text,
                    "content": [],
                    "level": level
                }
            else:
                current_section["content"].append(para.text)

        # Save last section
        if current_section["content"]:
            current_section["content"] = "\n".join(current_section["content"])
            doc.sections.append(current_section)

        # Extract tables
        for table in doc_file.tables:
            table_md = self._table_to_markdown(table)
            doc.tables.append(table_md)

        doc.content = "\n\n".join(s["heading"] + "\n" + s["content"] for s in doc.sections if s["content"])
        if doc.tables:
            doc.content += "\n\n--- Tables ---\n\n" + "\n\n".join(doc.tables)

        doc.metadata = {
            "paragraphs": len(doc_file.paragraphs),
            "tables": len(doc_file.tables),
            "sections": len(doc.sections),
        }
        doc._estimate_tokens()
        return doc

    def parse_html(self, content: bytes) -> ParsedDocument:
        """解析HTML文档"""
        self._check_size(content)  # 修复：读文件前检查大小上限
        if not self._bs4_available:
            return ParsedDocument(source_type="html", content="[HTML解析不可用: beautifulsoup4未安装]")

        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, "lxml")
        doc = ParsedDocument(source_type="html")

        # Remove scripts and styles
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        # Extract title
        doc.title = soup.title.string if soup.title else "Untitled HTML"

        # Extract headings and content
        for heading in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
            level = int(heading.name[1])
            text = heading.get_text(strip=True)
            # Find next sibling content
            content_parts = []
            for sibling in heading.find_next_siblings():
                if sibling.name and sibling.name.startswith("h"):
                    break
                content_parts.append(sibling.get_text(strip=True))

            if text:
                doc.sections.append({
                    "heading": text,
                    "content": " ".join(content_parts),
                    "level": level
                })

        # If no sections found, just extract all text
        if not doc.sections:
            doc.content = soup.get_text(separator="\n", strip=True)
        else:
            doc.content = "\n\n".join(
                "#" * s["level"] + " " + s["heading"] + "\n" + s["content"]
                for s in doc.sections
            )

        # Extract tables
        for table in soup.find_all("table"):
            table_md = self._html_table_to_markdown(table)
            doc.tables.append(table_md)

        doc._estimate_tokens()
        return doc

    def parse_text(self, content: bytes) -> ParsedDocument:
        """解析纯文本文件"""
        self._check_size(content)  # 修复：读文件前检查大小上限
        text = content.decode("utf-8", errors="replace")
        doc = ParsedDocument(source_type="txt", content=text)
        doc._estimate_tokens()
        return doc

    @staticmethod
    def _detect_extension(filename: str, content: bytes) -> str:
        """检测文件类型"""
        if filename:
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            if ext in ("pdf", "docx", "doc", "html", "htm", "txt"):
                return ext

        # Magic bytes detection
        if content[:4] == b"%PDF":
            return "pdf"
        elif content[:2] == b"PK":  # DOCX is a ZIP
            return "docx"
        elif b"<html" in content[:500].lower() or b"<!doctype html" in content[:500].lower():
            return "html"

        return "txt"

    @staticmethod
    def _table_to_markdown(table) -> str:
        """将python-docx表格转换为Markdown格式"""
        rows = []
        for row in table.rows:
            cells = [cell.text.replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) > 1:
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)

        return "\n".join(rows)

    @staticmethod
    def _html_table_to_markdown(table) -> str:
        """将HTML表格转换为Markdown格式"""
        from bs4 import BeautifulSoup  # Only import if available
        rows = []
        for tr in table.find_all("tr"):
            cells = [th.get_text(strip=True) for th in tr.find_all(["th", "td"])]
            if cells:
                rows.append("| " + " | ".join(cells) + " |")

        if len(rows) > 1:
            cell_count = len(rows[0].split("|")) - 2
            header_sep = "| " + " | ".join(["---"] * cell_count) + " |"
            rows.insert(1, header_sep)

        return "\n".join(rows)


# Add _estimate_tokens to ParsedDocument as a method
def _estimate_tokens(self):
    """估算token数量 (约4字符=1token for Chinese, ~4chars=1token for English)"""
    self.token_estimate = len(self.content) // 3  # Conservative estimate

ParsedDocument._estimate_tokens = _estimate_tokens
